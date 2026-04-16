from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def _add_runs_with_inline(p, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            p.add_run(part)


def md_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    lines = md_text.splitlines()
    in_code = False
    code_buf: list[str] = []

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            continue

        if in_code:
            code_buf.append(line)
            continue

        s = line.strip()
        if not s:
            continue

        if s == "---":
            p = doc.add_paragraph("—" * 24)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_after = Pt(6)
            continue

        if s.startswith("> "):
            p = doc.add_paragraph()
            _add_runs_with_inline(p, s[2:].strip())
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            continue

        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
            continue

        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_inline(p, s[2:].strip())
            continue

        p = doc.add_paragraph()
        _add_runs_with_inline(p, s)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert a simple Markdown file to .docx.")
    parser.add_argument("input_md", type=Path)
    parser.add_argument("-o", "--output", type=Path, default=None)
    args = parser.parse_args()

    md_path = args.input_md.resolve()
    if not md_path.exists():
        raise SystemExit(f"not found: {md_path}")

    out = args.output
    if out is None:
        out = md_path.with_suffix(".docx")
    else:
        out = out.resolve()

    md_to_docx(md_path.read_text(encoding="utf-8"), out)
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
